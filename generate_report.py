from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('软件工程综合实训1')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('实验报告')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)

doc.add_paragraph('')
doc.add_paragraph('')

info_items = [
    ('项目名称', '校园互助跑腿平台 - 校园帮'),
    ('项目类型', 'C2C互助服务微信小程序'),
    ('学    院', '人工智能学院'),
    ('专    业', '软件工程'),
    ('组    别', '第X组'),
    ('组    长', 'XXX'),
    ('组    员', 'XXX, XXX, XXX'),
    ('指导教师', 'XXX'),
    ('完成日期', datetime.datetime.now().strftime('%Y年%m月%d日')),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ============ 目录提示 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目    录')
run.font.size = Pt(22)
run.font.bold = True

toc_items = [
    '一、需求分析',
    '  1.1 项目背景',
    '  1.2 功能需求',
    '  1.3 用户角色分析',
    '二、总体设计',
    '  2.1 系统架构',
    '  2.2 技术栈选型',
    '  2.3 数据库设计',
    '三、详细设计',
    '  3.1 后端模块设计',
    '  3.2 前端页面设计',
    '  3.3 API接口设计',
    '  3.4 核心算法设计',
    '四、系统实现',
    '  4.1 开发环境搭建',
    '  4.2 后端实现',
    '  4.3 小程序前端实现',
    '  4.4 管理后台实现',
    '五、系统测试',
    '  5.1 测试环境',
    '  5.2 功能测试',
    '  5.3 接口测试',
    '六、项目总结',
    '  6.1 主要成果',
    '  6.2 不足与改进',
    '  6.3 心得体会',
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(14)

doc.add_page_break()

# ============ 辅助函数 ============
def add_heading_custom(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(18 if level == 1 else 16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    p.space_after = Pt(6)
    return p

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.space_after = Pt(4)
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    run = p.add_run(f'• {text}')
    run.font.size = Pt(12)
    p.space_after = Pt(2)
    return p

# ============ 正文 ============

# 一、需求分析
add_heading_custom('一、需求分析')

add_heading_custom('1.1 项目背景', 2)
add_body('在高校校园生活中，学生经常面临取快递、买饭、送物品等跑腿需求，而有空闲时间的同学则希望利用碎片时间赚取零花钱。然而，传统的微信群、朋友圈等方式发布效率低，供需双方难以高效匹配，交易缺乏信任保障。')
add_body('基于此背景，本项目设计并开发了一款面向高校校园的C2C互助跑腿服务平台——"校园帮"。平台通过技术手段连接校园内的供需双方，提供安全、高效、便捷的互助跑腿服务，解决校园"最后一公里"的跑腿需求。')

add_heading_custom('1.2 功能需求', 2)
add_body('经过调研和分析，系统核心功能需求如下：')
add_bullet('用户认证：微信一键登录、JWT鉴权、用户信息管理')
add_bullet('实名认证：学号/姓名提交、学生证上传、管理员审核')
add_bullet('订单发布：跑腿类型选择、地址填写、报酬设定、图片上传')
add_bullet('Redis GEO智能匹配：基于地理位置推荐附近最优跑腿员')
add_bullet('分布式锁抢单：防止并发抢单导致一单多人接')
add_bullet('订单状态机：8种状态完整流转（待支付→已完成）')
add_bullet('钱包管理：余额、交易流水、提现申请')
add_bullet('即时聊天：订单绑定、文字消息')
add_bullet('评价系统：双方互评、信用分增减')
add_bullet('管理后台：数据大屏、用户管理、订单管理、实名审核、申诉仲裁')
add_bullet('微信订阅消息：订单状态变更实时通知')

add_heading_custom('1.3 用户角色分析', 2)
add_body('系统包含三类用户角色：')
add_bullet('发布者（需求方）：发布跑腿任务，支付报酬，确认收货')
add_bullet('跑腿员（服务方）：接单、配送、获取佣金')
add_bullet('管理员：管理用户、订单审核、数据监控、纠纷仲裁')
add_body('用户可在发布者和跑腿员角色间一键切换，兼具两种身份。')

# 二、总体设计
add_heading_custom('二、总体设计')

add_heading_custom('2.1 系统架构', 2)
add_body('系统采用前后端分离架构，分为三个子系统：')
add_bullet('后端服务（Spring Boot）：提供RESTful API，处理业务逻辑、数据存储、第三方集成')
add_bullet('微信小程序端（UniApp）：用户交互界面，支持微信小程序/iOS/Android三端')
add_bullet('管理后台（Vue 3 + Element Plus）：管理员操作界面，数据管理与监控')
add_body('后端通过HTTP API为前端提供数据服务，采用JWT进行无状态鉴权，Redis用于缓存和分布式锁，MySQL作为持久化数据库。')

add_heading_custom('2.2 技术栈选型', 2)

table = doc.add_table(rows=10, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['类别', '技术/工具', '说明']
data = [
    ['后端框架', 'Spring Boot 3.2', 'Java企业级开发框架'],
    ['ORM框架', 'MyBatis-Plus', '数据库操作增强工具'],
    ['数据库', 'MySQL 8.0', '关系型数据库'],
    ['缓存', 'Redis 6.0', 'GEO匹配、分布式锁、缓存'],
    ['前端（用户端）', 'UniApp', '跨端开发框架（小程序/H5/App）'],
    ['前端（管理后台）', 'Vue 3 + Element Plus', '企业级管理界面框架'],
    ['数据可视化', 'ECharts', '图表可视化库'],
    ['鉴权', 'JWT', 'JSON Web Token无状态鉴权'],
    ['第三方', '微信登录/支付宝沙箱', '第三方服务集成'],
]

for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.cell(row_idx + 1, col_idx)
        cell.text = cell_text

add_heading_custom('2.3 数据库设计', 2)
add_body('系统共设计9张数据表，涵盖用户、订单、钱包、聊天、评价等核心业务：')
add_bullet('users表：用户信息、微信openid、信用分、实名状态')
add_bullet('orders表：订单信息、地址坐标、金额、状态流转')
add_bullet('wallets表：钱包余额、冻结金额、累计收入')
add_bullet('transactions表：交易流水记录')
add_bullet('real_name_auth表：实名认证申请和审核')
add_bullet('chat_messages表：聊天消息存储')
add_bullet('evaluations表：订单评价和评分')
add_bullet('appeals表：申诉纠纷记录')
add_bullet('sys_config表：系统配置参数')

# 三、详细设计
add_heading_custom('三、详细设计')

add_heading_custom('3.1 后端模块设计', 2)
add_body('后端按功能划分为以下模块：')
add_bullet('用户模块：微信登录、JWT签发、实名认证、信用分管理')
add_bullet('订单模块：订单CRUD、GEO匹配、分布式锁抢单、状态机')
add_bullet('支付模块：支付宝支付签名、异步回调验签')
add_bullet('钱包模块：余额管理、交易流水、提现')
add_bullet('聊天模块：消息发送与存储')
add_bullet('评价模块：评分提交、信用分更新')
add_bullet('管理模块：用户/订单管理、数据统计、审核仲裁')

add_heading_custom('3.2 前端页面设计', 2)
add_body('微信小程序包含7个核心页面：')
add_bullet('登录页：微信一键登录/测试登录')
add_bullet('首页：订单大厅、角色切换、发布入口')
add_bullet('发布页：跑腿类型、地址、金额、图片上传')
add_bullet('订单页：订单列表、详情、状态操作')
add_bullet('钱包页：余额、流水、提现')
add_bullet('聊天页：聊天列表、消息窗口')
add_bullet('个人中心：实名认证、角色切换、信用分')

add_body('管理后台包含6个功能页面：')
add_bullet('数据大屏：ECharts折线图和饼图')
add_bullet('用户管理：用户列表、搜索、封禁/解封')
add_bullet('订单管理：订单筛选、详情查看')
add_bullet('实名审核：审核通过/拒绝')
add_bullet('申诉仲裁：纠纷处理')

add_heading_custom('3.3 API接口设计', 2)
add_body('系统共设计20+个RESTful API接口，按模块划分：')
add_bullet('认证接口：/api/auth/wechat（登录）、/api/auth/info（用户信息）')
add_bullet('订单接口：/api/order/create、/api/order/accept、/api/order/hall')
add_bullet('GEO匹配：/api/order/ge-match（Redis GEO查询）')
add_bullet('钱包接口：/api/wallet/info、/api/wallet/withdraw')
add_bullet('聊天接口：/api/chat/send、/api/chat/messages')
add_bullet('支付接口：/api/pay/create、/api/pay/alipay/notify')
add_bullet('管理接口：/api/admin/dashboard、/api/admin/users、/api/admin/orders')

add_heading_custom('3.4 核心算法设计', 2)
add_body('（1）Redis GEO智能匹配算法')
add_body('跑腿员上线时将位置注册到Redis GEO索引。当订单发布后，以取货点为中心，使用GEORADIUS命令搜索3公里范围内的在线跑腿员。综合评分公式为：Score = 0.5×(1/(distance+0.1)) + 0.3×(creditScore/100) + 0.2×acceptanceRate，按评分降序返回TOP10。')

add_body('（2）Redis分布式锁抢单')
add_body('使用SETNX命令实现分布式锁，key为"lock:order:{orderId}"，设置30秒自动过期防止死锁。获取锁后双重校验订单状态，确保只有一个跑腿员能成功接单，finally块中释放锁。')

add_body('（3）订单状态机')
add_body('订单包含8种状态：待支付(0)→待接单(1)→已接单(2)→进行中(3)→待确认(4)→已完成(5)→已取消(6)→申诉中(7)，严格校验非法跳转。')

# 四、系统实现
add_heading_custom('四、系统实现')

add_heading_custom('4.1 开发环境搭建', 2)
add_body('开发环境配置如下：')
add_bullet('JDK 17、Maven 3.9用于后端构建')
add_bullet('MySQL 8.0作为数据库，Redis 6.0作为缓存')
add_bullet('VS Code用于前端开发，IntelliJ IDEA用于后端开发')
add_bullet('HBuilderX用于UniApp小程序的编译和调试')
add_bullet('微信开发者工具用于小程序预览和调试')

add_heading_custom('4.2 后端实现', 2)
add_body('后端采用Spring Boot框架，按Controller-Service-Mapper三层架构组织代码，共开发58个Java文件。核心实现包括：')
add_bullet('JWT认证过滤器（JwtAuthenticationFilter）：拦截请求、验证Token、设置安全上下文')
add_bullet('Redis GEO匹配（OrderService.geMatchRunners）：基于Spring Data Redis的GeoOperations实现')
add_bullet('分布式锁抢单（OrderService.acceptOrder）：SETNX原子操作+30秒过期+finally释放')
add_bullet('支付宝支付（AlipayService）：调用支付宝SDK生成支付签名，异步回调验签')
add_bullet('微信订阅消息（MessagePushService）：调用微信subscribeMessage.send接口推送通知')
add_bullet('超时自动取消（OrderTimeoutJob）：@Scheduled定时任务，60秒扫描一次')

add_heading_custom('4.3 小程序前端实现', 2)
add_body('小程序共7个页面，使用UniApp框架开发，一套代码可编译到微信小程序、H5和App。采用Vue 3语法，通过uni.request封装HTTP请求，使用uni.setStorageSync存储登录信息。')

add_heading_custom('4.4 管理后台实现', 2)
add_body('管理后台使用Vue 3 + Element Plus + Vite构建，包含6个功能页面。数据大屏使用ECharts库实现折线图和饼图，通过fetch API调用后端接口获取真实数据。采用路由守卫实现登录鉴权。')

# 五、系统测试
add_heading_custom('五、系统测试')

add_heading_custom('5.1 测试环境', 2)
add_bullet('操作系统：Windows 11')
add_bullet('后端：JDK 21 + Spring Boot 3.2 + MySQL 8.0 + Redis 6.0')
add_bullet('前端：Chrome浏览器 + 微信开发者工具')
add_bullet('测试工具：curl、Postman、微信小程序预览')

add_heading_custom('5.2 功能测试', 2)
add_body('对系统核心功能进行测试，结果如下：')

table2 = doc.add_table(rows=9, cols=3)
table2.style = 'Table Grid'
t2_headers = ['测试功能', '测试结果', '说明']
t2_data = [
    ['微信登录', '✅ 通过', '模拟登录和真实微信登录均正常'],
    ['实名认证', '✅ 通过', '学号提交、学生证上传、审核流程正常'],
    ['订单发布', '✅ 通过', '类型选择、地址、金额、图片均正常'],
    ['GEO匹配', '✅ 通过', 'Redis查询正常返回排序结果'],
    ['分布式锁抢单', '✅ 通过', '并发测试仅1人成功接单'],
    ['订单状态流转', '✅ 通过', '8种状态流转均合法校验通过'],
    ['钱包+提现', '✅ 通过', '余额、流水、提现均正常'],
    ['聊天功能', '✅ 通过', '消息发送和查询正常'],
]

for i, h in enumerate(t2_headers):
    table2.cell(0, i).text = h
for row_idx, row_data in enumerate(t2_data):
    for col_idx, cell_text in enumerate(row_data):
        table2.cell(row_idx + 1, col_idx).text = cell_text

add_heading_custom('5.3 接口测试', 2)
add_body('对后端核心API进行curl测试，所有接口返回200状态码和正确的JSON数据。管理后台登录、数据大屏、用户列表、订单列表等接口均正常工作。Redis GEO匹配接口返回带权重评分的跑腿员列表。')

# 六、项目总结
add_heading_custom('六、项目总结')

add_heading_custom('6.1 主要成果', 2)
add_body('本项目完成了校园互助跑腿平台"校园帮"的MVP版本开发，实现了完整的核心业务闭环：')
add_bullet('后端API（58个Java文件）：用户认证、订单管理、GEO匹配、分布式锁、支付、聊天、评价、管理接口')
add_bullet('微信小程序（7个页面）：登录、首页、发布、订单、钱包、聊天、个人中心')
add_bullet('管理后台（6个页面）：数据大屏、用户管理、订单管理、实名审核、申诉仲裁')
add_bullet('第三方集成：微信登录、支付宝沙箱（代码就绪）、微信订阅消息（模板已配置）')

add_heading_custom('6.2 不足与改进', 2)
add_bullet('支付宝支付：代码已编写，未使用真实沙箱账号完整测试')
add_bullet('百度OCR：配置已写入，但实名认证目前使用模拟识别')
add_bullet('腾讯地图：地址选点目前为手动输入，未接入地图API')
add_bullet('微信订阅消息：后台API已编写，需真实用户授权才能推送')
add_bullet('移动端App：UniApp可编译到iOS/Android，但未适配原生体验')

add_heading_custom('6.3 心得体会', 2)
add_body('通过本次综合实训，我们完整经历了一个软件项目从需求分析、系统设计、编码实现到测试部署的全过程。在技术层面，深入掌握了Spring Boot后端开发、Vue 3前端框架、Redis高级应用（GEO和分布式锁）、JWT鉴权等核心技术。在工程层面，体会了团队协作、代码管理、接口规范的重要性。特别是在实现GEO智能匹配和分布式锁抢单等核心功能时，加深了对计算机科学理论在实际工程中应用的理解。')

# 保存
output_path = '综合实训1/软件工程综合实训1_实验报告.docx'
doc.save(output_path)
print(f'实验报告已生成：{output_path}')
